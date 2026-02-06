import os
import subprocess
from pdf2docx import Converter
from PIL import Image

def pdf_to_docx(input_path, output_path):
    """Converts PDF to DOCX using pdf2docx."""
    try:
        cv = Converter(input_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
        return True
    except Exception as e:
        print(f"Error converting PDF to DOCX: {e}")
        return False

def libreoffice_convert(input_path, output_path, target_format):
    """Generic LibreOffice conversion."""
    try:
        out_dir = os.path.dirname(output_path)
        # Handle formats like 'pdf', 'docx', 'png', 'txt'
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', target_format, '--outdir', out_dir, input_path],
            capture_output=True,
            text=True
        )
        
        if result.returncode == 0:
            # LibreOffice output file name derivation
            basename = os.path.splitext(os.path.basename(input_path))[0]
            generated_file = os.path.join(out_dir, f"{basename}.{target_format}")
            
            if os.path.exists(generated_file):
                if generated_file != output_path:
                    if os.path.exists(output_path):
                        os.remove(output_path)
                    os.rename(generated_file, output_path)
                return True
        print(f"LibreOffice error: {result.stderr}")
        return False
    except Exception as e:
        print(f"Error in libreoffice_convert: {e}")
        return False

def docx_to_pdf(input_path, output_path):
    return libreoffice_convert(input_path, output_path, 'pdf')

def docx_to_excel(input_path, output_path):
    return libreoffice_convert(input_path, output_path, 'xlsx')

def docx_to_image(input_path, output_path):
    # Try LibreOffice
    return libreoffice_convert(input_path, output_path, 'png')

def excel_to_pdf(input_path, output_path):
    return libreoffice_convert(input_path, output_path, 'pdf')

def excel_to_word(input_path, output_path):
    return libreoffice_convert(input_path, output_path, 'docx')

def pdf_to_image(input_path, output_path):
    """Converts PDF to Image (first page)."""
    try:
        from pdf2image import convert_from_path
        images = convert_from_path(input_path, first_page=1, last_page=1)
        if images:
            images[0].save(output_path, 'PNG')
            return True
        return False
    except Exception as e:
        print(f"Error converting PDF to Image: {e}")
        return False

def image_to_word(input_path, output_path):
    """Extracts text from an image and saves it to a Word document."""
    try:
        from docx import Document
        import pytesseract
        from PIL import Image
        import shutil
        
        # Check if tesseract is installed
        if not shutil.which("tesseract"):
            print("Error: Tesseract OCR is not installed on the system.")
            return False
            
        # OCR extraction
        text = pytesseract.image_to_string(Image.open(input_path))
        
        doc = Document()
        if text.strip():
            doc.add_paragraph(text)
        else:
            # Fallback to embedding image if no text found
            from docx.shared import Inches
            doc.add_paragraph("No text could be extracted from the image. Embedded image below:")
            doc.add_picture(input_path, width=Inches(6))
            
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error in image_to_word: {e}")
        return False

def image_to_excel(input_path, output_path):
    """Extracts text from an image and saves it to an Excel spreadsheet."""
    try:
        import xlsxwriter
        import pytesseract
        from PIL import Image
        import shutil
        
        # Check if tesseract is installed
        if not shutil.which("tesseract"):
            print("Error: Tesseract OCR is not installed on the system.")
            return False

        # OCR extraction
        text = pytesseract.image_to_string(Image.open(input_path))
        lines = text.split('\n')
        
        workbook = xlsxwriter.Workbook(output_path)
        worksheet = workbook.add_worksheet()
        
        row = 0
        for line in lines:
            if line.strip():
                # Split line by multiple spaces or tabs to try and detect columns
                import re
                cols = re.split(r'\s{2,}|\t', line)
                for col_idx, col_val in enumerate(cols):
                    worksheet.write(row, col_idx, col_val.strip())
                row += 1
        
        if row == 0:
            # Fallback if no text extracted
            worksheet.write(0, 0, "No text extracted from image.")
            worksheet.insert_image('A2', input_path)
            
        workbook.close()
        return True
    except Exception as e:
        print(f"Error in image_to_excel: {e}")
        return False

def image_to_pdf(input_path, output_path):
    """Converts Image (JPG/PNG) to PDF using Pillow."""
    try:
        image = Image.open(input_path)
        if image.mode == 'RGBA':
            image = image.convert('RGB')
        image.save(output_path, "PDF", resolution=100.0)
        return True
    except Exception as e:
        print(f"Error converting Image to PDF: {e}")
        return False
